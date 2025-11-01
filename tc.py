import tkinter as tk
from tkinter import *
from tkinter import ttk
from tkinter.messagebox import showinfo, showerror
from connection import *
import screeninfo
from datetime import date
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches

today = date.today()

monitor = screeninfo.get_monitors()[0]

win = tk.Tk()
style = ttk.Style(win)
win.title('Рассчёты')
win.state('zoomed')
screen_height = monitor.height-110
screen_width = monitor.width

class FRAME:
    style = ttk.Style(win)

    def __init__(self):
        super(FRAME, self).__init__()

    def create_frame_left(self):
        self.left_frame = tk.Frame(win, width=f'{screen_width / 3}', height=f'{screen_height}',
                              pady=f'{screen_height / 32}', borderwidth=1, relief=SOLID, padx=f'{screen_width / 96}')
        self.right_frame = tk.Frame(win, borderwidth=1, relief=SOLID)
        self.left_frame.grid(row=0, column=0, sticky="nsew")
        self.right_frame.grid(row=0, column=1, sticky="nsew")

    def create_frame_top(self):
        self.top_frame = tk.Frame(win, width=f'{screen_width}', height=f'{screen_height/2}')
        self.bottom_frame = tk.Frame(win, pady=20)
        self.top_frame.grid(row=0, column=0, sticky="nsew")
        self.bottom_frame.grid(row=1, column=0, sticky="nsew")

class INGREDIENTS:
    style = ttk.Style(win)

    def __init__(self):
        super(INGREDIENTS, self).__init__()
        self.ingredients = []
        self.ingredient = ''
        self.weight_ingredient = 0
        self.cost = 0
        self.protein = 0
        self.fats = 0
        self.carbohydrates = 0
        self.calorie_content = 0

    def get_all_data_from_table_ingredients(self):
        select_data = """SELECT ingredient, weight, cost, protein, fats, carbohydrates, calorie_content, remainder FROM ingredients"""
        Data.cursor.execute(select_data)
        ingredient_data = Data.cursor.fetchall()
        for record in ingredient_data:
            self.ingredients.append(record)

    def change_data_into_table_ingredients(self, ingredient_entry: tk.Entry, weight_ingredient_entry: tk.Entry, cost_entry: tk.Entry,
                                           protein: tk.Entry, fats: tk.Entry, carbohydrates: tk.Entry, calorie_content: tk.Entry):
        ingredient_entry = ingredient_entry.get().upper()
        weight_ingredient_entry = weight_ingredient_entry.get()
        cost_entry = cost_entry.get()
        protein = protein.get()
        fats = fats.get()
        carbohydrates = carbohydrates.get()
        calorie_content = calorie_content.get()
        search_ingredient = """SELECT * FROM ingredients WHERE ingredient LIKE ?"""
        Data.cursor.execute(search_ingredient, (ingredient_entry,))
        get_data = Data.cursor.fetchone()
        if get_data != None:
            remainder = get_data[8] + (float(weight_ingredient_entry)/1000)
            ingredients_update = (weight_ingredient_entry, cost_entry, remainder, protein, fats, carbohydrates, calorie_content, ingredient_entry)
            update_data_into_table_ingredients = """
    				UPDATE ingredients SET weight = ?, cost = ?, remainder = ?, protein = ?, fats = ?, carbohydrates = ?, calorie_content = ? WHERE ingredient = ?"""
            Data.cursor.execute(update_data_into_table_ingredients, ingredients_update)
        else:
            ingredients_entry = (ingredient_entry, weight_ingredient_entry, cost_entry, (int(weight_ingredient_entry)/1000), protein, fats, carbohydrates, calorie_content)
            insert_data_into_table_ingredients = """
    				    INSERT INTO ingredients (ingredient, weight, cost, remainder, protein, fats, carbohydrates, calorie_content)
    				    VALUES (?, ?, ?, ?, ?, ?, ?, ?);"""
            Data.cursor.execute(insert_data_into_table_ingredients, ingredients_entry)
        Data.db.commit()
        self.add_ingredients()

    def show_table_ingredients(self):
        F.create_frame_left()
        INGREDIENTS.style.theme_use("clam")
        INGREDIENTS.style.configure('Treeview', background='grey', foreground='white',
                                    fieldbackground='grey', sticky='nsew')
        heads = ['ИНГРЕДИЕНТ', 'ВЕС/КОЛИЧЕСТВО, Г/МЛ', 'СТОИМОСТЬ', 'БЕЛКИ', 'ЖИРЫ', 'УГЛЕВОДЫ', 'КАЛОРИЙНОСТЬ', 'ОСТАТОК, КГ/Л']
        table = ttk.Treeview(F.right_frame, show='headings')
        table['columns'] = heads
        for header in heads:
            table.heading(header, text=header, anchor='center')
            table.column(header, anchor='center')
        for row in self.ingredients:
            table.insert('', tk.END, values=row)
        table.column(0, stretch=False, width=int(f"{int(screen_width / 11.8)}"))
        table.column(1, stretch=False, width=int(f"{int(screen_width / 12.6)}"))
        table.column(2, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        table.column(3, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        table.column(4, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        table.column(5, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        table.column(6, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        table.column(7, stretch=False, width=int(f"{int(screen_width / 12.9)}"))
        scroll_pane = ttk.Scrollbar(F.right_frame, command=table.yview)
        table.configure(yscrollcommand=scroll_pane.set)
        scroll_pane.pack(side=tk.RIGHT, fill=tk.BOTH)
        table.pack(expand=tk.YES, fill=tk.BOTH)
        self.ingredients = []
        def selected_ingredient(event):
            for selected_item in table.selection():
                item = table.item(selected_item)
                ingredient_data = item["values"]
                self.ingredient = ingredient_data[0]
                self.weight_ingredient = ingredient_data[1]
                self.cost = ingredient_data[2]
                self.protein = ingredient_data[3]
                self.fats = ingredient_data[4]
                self.carbohydrates = ingredient_data[5]
                self.calorie_content = ingredient_data[6]
            self.add_ingredients()
        table.bind("<<TreeviewSelect>>", selected_ingredient)

    def add_ingredients(self):
        [child.destroy() for child in win.winfo_children()]
        create_menubar()
        F.create_frame_left()
        self.get_all_data_from_table_ingredients()
        self.show_table_ingredients()
        tk.Label(F.left_frame, text='Ингредиент', font=f'Arial {int(screen_height/45)} italic underline').grid(row=0, column=0, sticky='w')
        ingredient_entry = tk.Entry(F.left_frame, font=f'Arial {int(screen_height/45)}')
        ingredient_entry.grid(row=0, column=1, padx=5, pady=5)
        ingredient_entry.insert(0, self.ingredient)
        ingredient_entry.focus()
        tk.Label(F.left_frame, text='Количество, гр', font=f'Arial {int(screen_height/45)} italic underline').grid(row=1, column=0, sticky='w')
        weight_ingredient_entry = tk.Entry(F.left_frame, font=f'Arial {int(screen_height/45)}')
        weight_ingredient_entry.grid(row=1, column=1, padx=5, pady=5)
        weight_ingredient_entry.insert(0, self.weight_ingredient)
        tk.Label(F.left_frame, text='Стоимость', font=f'Arial {int(screen_height/45)} italic underline').grid(row=2, column=0, sticky='w')
        cost_entry = tk.Entry(F.left_frame, font=f'Arial {int(screen_height/45)}')
        cost_entry.grid(row=2, column=1, padx=5, pady=5)
        cost_entry.insert(0, self.cost)
        tk.Label(F.left_frame, text='Белки', font=f'Arial {int(screen_height /45)} italic underline').grid(row=3,
                                                                                                                column=0,
                                                                                                                sticky='w')
        protein = tk.Entry(F.left_frame, font=f'Arial {int(screen_height /45)}')
        protein.grid(row=3, column=1, padx=5, pady=5)
        protein.insert(0, self.protein)
        tk.Label(F.left_frame, text='Жиры', font=f'Arial {int(screen_height / 45)} italic underline').grid(row=4,
                                                                                                                column=0,
                                                                                                                sticky='w')
        fats = tk.Entry(F.left_frame, font=f'Arial {int(screen_height / 45)}')
        fats.grid(row=4, column=1, padx=5, pady=5)
        fats.insert(0, self.fats)
        tk.Label(F.left_frame, text='Углеводы', font=f'Arial {int(screen_height / 45)} italic underline').grid(row=5,
                                                                                                                column=0,
                                                                                                                sticky='w')
        carbohydrates = tk.Entry(F.left_frame, font=f'Arial {int(screen_height / 45)}')
        carbohydrates.grid(row=5, column=1, padx=5, pady=5)
        carbohydrates.insert(0, self.carbohydrates)
        tk.Label(F.left_frame, text='Калорийность', font=f'Arial {int(screen_height / 45)} italic underline').grid(row=6,
                                                                                                                column=0,
                                                                                                                sticky='w')
        calorie_content = tk.Entry(F.left_frame, font=f'Arial {int(screen_height / 45)}')
        calorie_content.grid(row=6, column=1, padx=5, pady=5)
        calorie_content.insert(0, self.calorie_content)
        btn = tk.Button(F.left_frame, text='ДОБАВИТЬ / ИЗМЕНИТЬ', font=f'Arial {int(screen_height/43)} bold',
                        command=lambda: self.change_data_into_table_ingredients(ingredient_entry, weight_ingredient_entry,
                                                                                cost_entry, protein, fats, carbohydrates, calorie_content))
        btn.grid(row=7, column=0, padx=5, columnspan=2)
        self.ingredient = ''
        self.weight_ingredient = 0
        self.cost = 0
        self.protein = 0
        self.fats = 0
        self.carbohydrates = 0
        self.calorie_content = 0


class PRODUCTS:
    style = ttk.Style(win)

    def __init__(self):
        super(PRODUCTS, self).__init__()
        self.products_description = []
        self.products_names = []
        self.composition = []
        self.product_name = ''
        self.appearance = ''
        self.consistency = ''
        self.taste = ''
        self.smell = ''

    # ОПИСАНИЕ ПРОДУКТА
    def get_all_data_from_table_products_description(self, type):
        select_data = """SELECT product_name, appearance, consistency, taste, smell FROM products_description"""
        Data.cursor.execute(select_data)
        product_description_data = Data.cursor.fetchall()
        if type == 'description':
            for record in product_description_data:
                self.products_description.append(record)
        if type == 'names':
            for record in product_description_data:
                self.products_names.append(record[0])
    def change_data_into_table_products_description(self, product_name: tk.Entry, appearance: tk.Entry, consistency: tk.Entry, taste: tk.Entry, smell: tk.Entry):
        product_name = product_name.get().upper()
        appearance = appearance.get()
        consistency = consistency.get()
        taste = taste.get()
        smell = smell.get()
        search_product_description = """SELECT * FROM products_description WHERE product_name LIKE ?"""
        Data.cursor.execute(search_product_description, (product_name,))
        get_data = Data.cursor.fetchone()
        if get_data != None:
            product_description_update = (appearance, consistency, taste, smell, product_name)
            update_data_into_table_products_description = """
    				UPDATE products_description SET appearance = ?, consistency = ?, taste = ?, smell=? WHERE product_name = ?"""
            Data.cursor.execute(update_data_into_table_products_description, product_description_update)
        else:
            product_description_entry = (product_name, appearance, consistency, taste, smell)
            insert_data_into_table_products_description = """
    				    INSERT INTO products_description (product_name, appearance, consistency, taste, smell)
    				    VALUES (?, ?, ?, ?, ?);"""
            Data.cursor.execute(insert_data_into_table_products_description, product_description_entry)
        Data.db.commit()
        self.product_description()

    def show_table_products_description(self):
        F.create_frame_top()
        PRODUCTS.style.theme_use("clam")
        PRODUCTS.style.configure('Treeview', background='grey', foreground='white',
                                    fieldbackground='grey', sticky='nw')
        heads = ['ПРОДУКТ', 'ВНЕШНИЙ ВИД', 'КОНСИСТЕНЦИЯ', 'ВКУС', 'ЗАПАХ']
        table = ttk.Treeview(F.top_frame, show='headings')
        table['columns'] = heads
        for header in heads:
            table.heading(header, text=header, anchor='center')
            table.column(header, anchor='w')
        for row in self.products_description:
            table.insert('', tk.END, values=row)
        table.column(0, stretch=False, width=f"{int(screen_width / 6)}")
        table.column(1, stretch=False, width=f"{int(screen_width / 3.85)}")
        table.column(2, stretch=False, width=f"{int(screen_width / 5.3)}")
        table.column(3, stretch=False, width=f"{int(screen_width / 5.3)}")
        table.column(4, stretch=False, width=f"{int(screen_width / 5.3)}")
        scroll_pane = ttk.Scrollbar(F.top_frame, command=table.yview)
        table.configure(yscrollcommand=scroll_pane.set)
        scroll_pane.pack(side=tk.RIGHT, fill=tk.BOTH)
        table.pack(expand=tk.YES, fill=tk.BOTH)
        self.products_description = []
        def selected_product(event):
            for selected_item in table.selection():
                item = table.item(selected_item)
                product_data = item["values"]
                self.product_name = product_data[0]
                self.appearance = product_data[1]
                self.consistency = product_data[2]
                self.taste = product_data[3]
                self.smell = product_data[4]
            self.product_description()
        table.bind("<<TreeviewSelect>>", selected_product)

    def product_description(self):
        [child.destroy() for child in win.winfo_children()]
        create_menubar()
        F.create_frame_top()
        self.get_all_data_from_table_products_description('description')
        self.show_table_products_description()
        tk.Label(F.bottom_frame, text='Продукт:', font=f'Arial {int(screen_height/32)} italic').grid(row=0, column=0, sticky='w')
        product_name = tk.Entry(F.bottom_frame, font=f'Arial {int(screen_height/32)}')
        product_name.grid(row=0, column=1, padx=25, pady=25)
        product_name.insert(0, self.product_name)
        product_name.focus()
        tk.Label(F.bottom_frame, text='Внешний вид:', font=f'Arial {int(screen_height/32)} italic').grid(row=0, column=2, sticky='w')
        appearance = tk.Entry(F.bottom_frame, font=f'Arial {int(screen_height/32)}')
        appearance.grid(row=0, column=3, padx=25, pady=25)
        appearance.insert(0, self.appearance)
        tk.Label(F.bottom_frame, text='Консистенция:', font=f'Arial {int(screen_height/32)} italic').grid(row=1, column=0, sticky='w')
        consistency = tk.Entry(F.bottom_frame, font=f'Arial {int(screen_height/32)}')
        consistency.grid(row=1, column=1, padx=25, pady=25)
        consistency.insert(0, self.consistency)
        tk.Label(F.bottom_frame, text='Вкус:', font=f'Arial {int(screen_height / 32)} italic').grid(
            row=1, column=2, sticky='w')
        taste = tk.Entry(F.bottom_frame, font=f'Arial {int(screen_height / 32)}')
        taste.grid(row=1, column=3, padx=25, pady=25)
        taste.insert(0, self.taste)
        tk.Label(F.bottom_frame, text='Запах:', font=f'Arial {int(screen_height / 32)} italic').grid(
            row=2, column=0, sticky='w')
        smell = tk.Entry(F.bottom_frame, font=f'Arial {int(screen_height / 32)}')
        smell.grid(row=2, column=1, padx=25, pady=25)
        smell.insert(0, self.smell)
        btn = tk.Button(F.bottom_frame, text='ДОБАВИТЬ / ИЗМЕНИТЬ', font=f'Arial {int(screen_height/32)} bold',
                        command=lambda: self.change_data_into_table_products_description(product_name, appearance, consistency, taste, smell))
        btn.grid(row=3, column=1,)
        self.product_name = ''
        self.appearance = ''
        self.consistency = ''
        self.taste = ''
        self.smell = ''

    # СОСТАВ ПРОДУКТА
    def get_ingredients_from_table_ingredients(self):
        select_ingredients = """SELECT ingredient FROM ingredients"""
        Data.cursor.execute(select_ingredients)
        ingredients = Data.cursor.fetchall()
        return ingredients

    def get_product_composition(self, product_name: tk.Entry, type):
        product_name = product_name.get()
        select_product_composition = """SELECT product_name, ingredient, quantity FROM products_composition WHERE product_name LIKE ?"""
        Data.cursor.execute(select_product_composition, (product_name,))
        product_composition = Data.cursor.fetchall()
        self.composition = product_composition
        if type == 'calculation_card':
            C.price_calculation()
        else:
            self.product_composition()

    def change_data_into_table_products_composition(self, product_name: tk.Entry, ingredient: tk.Entry, quantity: tk.Entry):
        product_name = product_name.get()
        ingredient = ingredient.get()
        quantity = quantity.get()
        search_product_composition = """SELECT quantity FROM products_composition WHERE product_name LIKE ? and ingredient LIKE ?"""
        Data.cursor.execute(search_product_composition, (product_name, ingredient,))
        get_data = Data.cursor.fetchone()
        if get_data != None and get_data != quantity:
            product_composition_update = (quantity, product_name, ingredient)
            update_data_into_table_products_composition = """
            				UPDATE products_composition SET quantity = ? WHERE product_name = ? and ingredient = ?"""
            Data.cursor.execute(update_data_into_table_products_composition, product_composition_update)
        elif get_data == None:
            product_composition_entry = (product_name, ingredient, quantity)
            insert_data_into_table_products_composition = """
            				    INSERT INTO products_composition (product_name, ingredient, quantity)
            				    VALUES (?, ?, ?);"""
            Data.cursor.execute(insert_data_into_table_products_composition, product_composition_entry)
        Data.db.commit()
        select_product_composition = """SELECT product_name, ingredient, quantity FROM products_composition WHERE product_name LIKE ?"""
        Data.cursor.execute(select_product_composition, (product_name,))
        product_composition = Data.cursor.fetchall()
        self.composition = product_composition
        self.product_composition()

    def show_table_product_composition(self):
        F.create_frame_left()
        PRODUCTS.style.theme_use("clam")
        PRODUCTS.style.configure('Treeview', background='grey', foreground='white',
                                    fieldbackground='grey', sticky='nw')
        heads = ['ПРОДУКТ', 'ИНГРЕДИЕНТ', 'КОЛИЧЕСТВО']
        table = ttk.Treeview(F.right_frame, show='headings')
        table['columns'] = heads
        for header in heads:
            table.heading(header, text=header, anchor='center')
            table.column(header, anchor='w')
        for row in self.composition:
            table.insert('', tk.END, values=row)
        table.column(0, stretch=False, width=f"{int(screen_width / 5.3)}")
        table.column(1, stretch=False, width=f"{int(screen_width / 5.3)}")
        table.column(2, stretch=False, width=f"{int(screen_width / 5.3)}")
        scroll_pane = ttk.Scrollbar(F.right_frame, command=table.yview)
        table.configure(yscrollcommand=scroll_pane.set)
        scroll_pane.pack(side=tk.RIGHT, fill=tk.BOTH)
        table.pack(expand=tk.YES, fill=tk.BOTH)
        self.composition = []


    def product_composition(self):
        [child.destroy() for child in win.winfo_children()]
        create_menubar()
        F.create_frame_left()
        self.get_all_data_from_table_products_description('names')
        self.show_table_product_composition()
        tk.Label(F.left_frame, text='Продукт:', font=f'Arial {int(screen_height/37)} italic').grid(row=0, column=0, sticky='w')
        product_name_combobox = ttk.Combobox(F.left_frame, state='readonly', font=f'Arial {int(screen_height/37)} italic')
        product_name_combobox['values'] = self.products_names
        product_name_combobox.grid(row=0, column=1, pady=10)
        product_name_combobox.current(0)
        product_name = StringVar(value=product_name_combobox.get())
        def selected_product(event):
            product_name = StringVar(value=product_name_combobox.get())
            selected_product_name['text'] = product_name
        product_name_combobox.bind("<<ComboboxSelected>>", selected_product)
        selected_product_name = tk.Entry(textvariable=product_name)
        product_name_combobox.focus()
        show_btn = tk.Button(F.left_frame, text='Cостав продукта', font=f'Arial {int(screen_height / 37)} bold',
                            command=lambda: self.get_product_composition(product_name_combobox, 'product_composition'))
        show_btn.grid(row=1, column=1)
        tk.Label(F.left_frame, text='Ингредиент:', font=f'Arial {int(screen_height/37)} italic').grid(row=2, column=0, sticky='w')
        ingredients_combobox = ttk.Combobox(F.left_frame, state='readonly',
                                             font=f'Arial {int(screen_height / 37)} italic')
        ingredients_combobox['values'] = self.get_ingredients_from_table_ingredients()
        ingredients_combobox.grid(row=2, column=1, pady=10)
        ingredients_combobox.current(0)
        ingredient = StringVar(value=ingredients_combobox.get())
        def selected_ingredient(event):
            ingredient = StringVar(value=ingredients_combobox.get())
            selected_ingredient['text'] = ingredient
        ingredients_combobox.bind("<<ComboboxSelected>>", selected_ingredient)
        selected_ingredient = tk.Entry(textvariable=ingredient)
        tk.Label(F.left_frame, text='Количество,\n гр/мл:', font=f'Arial {int(screen_height/37)} italic').grid(row=3, column=0, sticky='w')
        quantity = tk.Entry(F.left_frame, font=f'Arial {int(screen_height/37)}')
        quantity.grid(row=3, column=1, pady=25)
        add_btn = tk.Button(F.left_frame, text='ДОБАВИТЬ / ИЗМЕНИТЬ', font=f'Arial {int(screen_height/37)} bold',
                        command=lambda: self.change_data_into_table_products_composition(product_name_combobox, ingredients_combobox, quantity))
        add_btn.grid(row=4, column=0, columnspan=2)
        self.products_names = []

class CARDS:
    style = ttk.Style(win)

    def __init__(self):
        super(CARDS, self).__init__()
        self.product = ''
        self.product_weight = 0
        self.price_ingredients = []
        self.product_price = round(0, 2)

    def price_calculation(self):
        self.product = P.composition[0][0]
        for elem in P.composition:
            select_ingredient = """SELECT ingredient, weight, cost FROM ingredients WHERE ingredient LIKE ?"""
            Data.cursor.execute(select_ingredient, (elem[1],))
            ingredient_data = Data.cursor.fetchall()
            self.product_weight += elem[2]
            price = round((elem[2] * ingredient_data[0][2]) / ingredient_data[0][1], 2)
            self.product_price += price
            ingredient_price = elem[1], elem[2], price
            self.price_ingredients.append(ingredient_price)
        self.create_a_printed_version()
        self.create_calculation_card()

    def show_calculation_card(self):
        F.create_frame_left()
        tk.Label(F.right_frame, text='Калькуляционная карта:', font=f'Arial {int(screen_height / 37)} italic').grid(row=0, column=1)
        tk.Label(F.right_frame, text='от ' f'{today.strftime("%d.%m.%Y")}', font=f'Arial {int(screen_height / 37)} italic').grid(
            row=1, column=1)
        tk.Label(F.right_frame, text=f'{self.product}', font=f'Arial {int(screen_height / 37)} italic').grid(
            row=2, column=1)
        tree = ttk.Treeview(F.right_frame, columns=['Ингредиент', 'Количество', 'Стоимость'], show="headings")
        style = ttk.Style()
        style.configure("Treeview.Heading", font=(None, 20))
        heads = ['Ингредиент', 'Количество', 'Стоимость']
        tree['columns'] = heads
        for header in heads:
            tree.heading(header, text=header, anchor='center')
            tree.column(header, anchor='center')
        for row in self.price_ingredients:
            tree.insert('', tk.END, values=row)
        tree.column(0, stretch=False, width=f"{int(screen_width / 5.2)}")
        tree.column(1, stretch=False, width=f"{int(screen_width / 5.2)}")
        tree.column(2, stretch=False, width=f"{int(screen_width / 5.2)}")
        tree.grid(row=3, column=0, columnspan=3, padx=20, pady=20)
        tk.Label(F.right_frame, text=f'Вес продукта: {round(self.product_weight, 2)/1000}кг',
                 font=f'Arial {int(screen_height / 37)} italic').grid(row=4, column=0, columnspan=2)
        tk.Label(F.right_frame, text=f'Стоимость продукта: {round(self.product_price, 2)}руб', font=f'Arial {int(screen_height / 37)} italic').grid(row=5, column=0, columnspan=2)

    def create_a_printed_version(self):
            doc = docx.Document()
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            section = doc.sections[0]
            section.different_first_page_header_footer = True
            first_page_header = section.first_page_header
            first_page_header = False
            style.paragraph_format.right_indent = Inches(-0.5)
            style.paragraph_format.left_indent = Inches(-0.5)
            style.paragraph_format.space_after = Inches(0)
            par_1 = doc.add_paragraph()
            run_1 = par_1.add_run('КАЛЬКУЛЯЦИОННАЯ КАРТА\n')
            run_1.font.bold = True
            run_1.font.size = Pt(14)
            par_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_2 = par_1.add_run('от ' f'{today.strftime("%d.%m.%Y")}\n')
            run_2.font.bold = True
            run_3 = par_1.add_run(f'{self.product}\n')
            run_3.font.size = Pt(11)
            run_3.font.bold = True
            header_ingredients = ['ИНГРЕДИЕНТ', 'КОЛИЧЕСТВО', 'СТОИМОСТЬ']
            table_ingredients = doc.add_table(1, len(header_ingredients), style='Table Grid')
            table_ingredients.alignment = WD_TABLE_ALIGNMENT.CENTER
            head_cells = table_ingredients.rows[0].cells
            for i, item in enumerate(header_ingredients):
                p = head_cells[i].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                head_row = p.add_run(item)
                head_row.font.bold = True
            for row in self.price_ingredients:
                cells = table_ingredients.add_row().cells
                for i, item in enumerate(row):
                    cells[i].text = str(item)
                for cell in cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for cell in table_ingredients.columns[0].cells:
                cell.width = Inches(3.0)
            for row in table_ingredients.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size = Pt(12)
            weight = doc.add_paragraph()
            run_weight = weight.add_run(f'Вес продукта: {self.product_weight}')
            run_weight.font.bold = True
            cost = doc.add_paragraph()
            run_cost = cost.add_run(f'Стоимость продукта: {round(self.product_price, 2)}')
            doc.save(f'КАЛЬКУЛЯЦИОННЫЕ КАРТЫ/{self.product} от {today}.docx')
            self.price_ingredients = []


    def create_calculation_card(self):
        [child.destroy() for child in win.winfo_children()]
        create_menubar()
        F.create_frame_left()
        self.show_calculation_card()
        P.get_all_data_from_table_products_description('names')
        tk.Label(F.left_frame, text='Продукт:', font=f'Arial {int(screen_height / 37)} italic').grid(row=0, column=0,
                                                                                                        sticky='w')
        product_name_combobox = ttk.Combobox(F.left_frame, state='readonly',
                                             font=f'Arial {int(screen_height / 37)} italic')
        product_name_combobox['values'] = P.products_names
        product_name_combobox.grid(row=0, column=1, pady=10)
        product_name_combobox.current(0)
        product_name = StringVar(value=product_name_combobox.get())
        def selected_product(event):
            product_name = StringVar(value=product_name_combobox.get())
            selected_product_name['text'] = product_name
        product_name_combobox.bind("<<ComboboxSelected>>", selected_product)
        selected_product_name = tk.Entry(textvariable=product_name)
        product_name_combobox.focus()
        show_btn = tk.Button(F.left_frame, text='РАССЧИТАТЬ', font=f'Arial {int(screen_height / 37)} bold',
                             command=lambda: P.get_product_composition(product_name_combobox, 'calculation_card'))
        show_btn.grid(row=1, column=1, pady=10)
        P.products_names = []
        P.composition = []
        self.product = ''
        self.price_ingredients = []
        self.product_price = round(0, 2)
        self.product_weight = 0
def create_menubar():
    menubar = tk.Menu(win)
    win.config(menu=menubar)
    directory = tk.Menu(menubar, tearoff=0)
    directory.add_command(label='Ингредиенты', command=I.add_ingredients)
    directory.add_command(label='Описание продукта', command=P.product_description)
    directory.add_command(label='Cостав продукта', command=P.product_composition)
    menubar.add_cascade(label='Справочники', menu=directory)
    cards = tk.Menu(menubar, tearoff=0)
    cards.add_command(label='Калькуляционная карта', command=C.create_calculation_card)
    menubar.add_cascade(label='Карты', menu=cards)

def start():
    create_menubar()
    win.mainloop()

F = FRAME()
I = INGREDIENTS()
P = PRODUCTS()
C = CARDS()
programm = start()
